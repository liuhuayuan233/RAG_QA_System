import os
import logging
import json
from typing import List, Dict, Any, Optional
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LangchainDocument

from config.config import Config
from src.utils import clean_text, validate_file_size, get_file_extension, setup_logging

logger = setup_logging()

class DocumentProcessor:
    """文档处理器"""
    
    def __init__(self):
        self.config = Config()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.CHUNK_SIZE,
            chunk_overlap=self.config.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""]
        )
    
    def process_document(self, file_path: str) -> List[LangchainDocument]:
        """处理单个文档"""
        try:
            # 验证文件
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")
            
            if not validate_file_size(file_path, self.config.MAX_DOCUMENT_SIZE):
                raise ValueError(f"文件过大: {file_path}")
            
            # 获取文件扩展名
            ext = get_file_extension(file_path)
            if ext not in self.config.SUPPORTED_EXTENSIONS:
                raise ValueError(f"不支持的文件格式: {ext}")
            
            # 提取文本
            text = self._extract_text(file_path, ext)
            if not text or len(text.strip()) < 50:
                raise ValueError(f"文档内容过短或为空: {file_path}")
            
            # 清理文本
            text = clean_text(text)
            
            # 分块处理
            chunks = self.text_splitter.split_text(text)
            
            # 创建Document对象
            documents = []
            filename = Path(file_path).name
            
            for i, chunk in enumerate(chunks):
                if len(chunk.strip()) < 20:  # 过滤太短的块
                    continue
                    
                doc = LangchainDocument(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "filename": filename,
                        "chunk_id": i,
                        "total_chunks": len(chunks),
                        "file_type": self.config.SUPPORTED_EXTENSIONS[ext]
                    }
                )
                documents.append(doc)
            
            logger.info(f"成功处理文档 {filename}，生成 {len(documents)} 个文档块")
            return documents
            
        except Exception as e:
            logger.error(f"处理文档失败 {file_path}: {str(e)}")
            raise
    
    def _extract_text(self, file_path: str, ext: str) -> str:
        """根据文件类型提取文本"""
        if ext == ".pdf":
            return self._extract_pdf_text(file_path)
        elif ext == ".docx":
            return self._extract_docx_text(file_path)
        elif ext in [".txt", ".md"]:
            return self._extract_text_file(file_path)
        elif ext == ".jsonl":
            return self._extract_jsonl_text(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """提取PDF文本"""
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()
                if page_text.strip():
                    text += f"\n[第{page_num + 1}页]\n{page_text}\n"
            
            doc.close()
            return text
            
        except Exception as e:
            logger.error(f"提取PDF文本失败 {file_path}: {str(e)}")
            raise
    
    def _extract_docx_text(self, file_path: str) -> str:
        """提取Word文档文本"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # 处理表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        text += row_text + "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"提取Word文档文本失败 {file_path}: {str(e)}")
            raise
    
    def _extract_text_file(self, file_path: str) -> str:
        """提取文本文件内容"""
        try:
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            raise ValueError(f"无法解码文件 {file_path}")
            
        except Exception as e:
            logger.error(f"提取文本文件失败 {file_path}: {str(e)}")
            raise
    
    def process_directory(self, directory_path: str) -> List[LangchainDocument]:
        """处理目录中的所有文档"""
        try:
            if not os.path.exists(directory_path):
                raise FileNotFoundError(f"目录不存在: {directory_path}")
            
            all_documents = []
            processed_files = 0
            
            for root, dirs, files in os.walk(directory_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    ext = get_file_extension(file_path)
                    
                    if ext in self.config.SUPPORTED_EXTENSIONS:
                        try:
                            documents = self.process_document(file_path)
                            all_documents.extend(documents)
                            processed_files += 1
                            
                        except Exception as e:
                            logger.warning(f"跳过文件 {file_path}: {str(e)}")
                            continue
            
            logger.info(f"成功处理 {processed_files} 个文件，生成 {len(all_documents)} 个文档块")
            return all_documents
            
        except Exception as e:
            logger.error(f"处理目录失败 {directory_path}: {str(e)}")
            raise 

    def _extract_jsonl_text(self, file_path: str) -> str:
        """提取JSONL文件中的医疗问答文本"""
        try:
            texts = []
            
            with open(file_path, 'r', encoding='utf-8') as f:
                for line_num, line in enumerate(f, 1):
                    line = line.strip()
                    if not line:
                        continue
                    
                    try:
                        # 解析JSON行
                        item = json.loads(line)
                        
                        # 提取问题和答案
                        questions = item.get('questions', [])
                        answers = item.get('answers', [])
                        
                        if questions and answers:
                            # 处理问题（可能是嵌套数组）
                            if isinstance(questions[0], list) and len(questions[0]) > 0:
                                question = questions[0][0].strip()
                            else:
                                question = str(questions[0]).strip() if questions[0] else ""
                            
                            # 处理答案
                            answer = answers[0].strip() if len(answers) > 0 else ""
                            
                            if question and answer:
                                # 格式化为问答对
                                qa_text = f"问题：{question}\n\n答案：{answer}"
                                texts.append(qa_text)
                    
                    except json.JSONDecodeError:
                        logger.warning(f"跳过无效JSON行 {line_num}: {file_path}")
                        continue
                    except Exception as e:
                        logger.warning(f"处理行 {line_num} 时出错: {e}")
                        continue
            
            combined_text = "\n\n" + "="*50 + "\n\n".join(texts)
            logger.info(f"JSONL文件处理完成，提取了 {len(texts)} 个问答对")
            
            return combined_text
            
        except Exception as e:
            logger.error(f"读取JSONL文件失败 {file_path}: {str(e)}")
            raise
